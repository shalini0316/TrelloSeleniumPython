


from selenium import webdriver
from docx import Document
import docx
from docx.shared import Inches

ss_fileName = []
class SS(object):

    def __init__(self, driver):
        self.driver = driver

    def ScreenShot(self, path):
        directory = "C:\\Users\\BLACKPEARL COMPUTERS\\PycharmProjects\\pythonProject\\PythonSelFramework\\tests\\screenshot\\"
        ss_fileName.append(directory+path)

        self.driver.get_screenshot_as_file(directory+path)


    def writeDoc(self,name):

        mydoc = docx.Document()
        mydoc.add_heading(name, 0)
        mydoc.add_paragraph('1.Login into Trello.')
        mydoc.add_paragraph('2.Create a new Board. (Project)')
        mydoc.add_paragraph('3.Create 4 Lists (Not Started, In Progress, QA, Done )')
        mydoc.add_paragraph('4.Create 4 Cards. : Card 1, Card 2, Card 3, Card 4. under the list Not Started.')
        mydoc.add_paragraph('5.Move Card 2 to In Progress.')
        mydoc.add_paragraph('6.Move Card 3 to QA.')
        mydoc.add_paragraph('7.Move Card 2 under QA.')
        mydoc.add_paragraph('8.Open Card 1 and Assign it to the current logged in user and then leave a comment on Card 1 saying “I am done”')

        for img in ss_fileName:
            print(img)

            mydoc.add_picture(img, width=Inches(7))
        mydoc.save("C:\\Users\\BLACKPEARL COMPUTERS\\PycharmProjects\\pythonProject\\PythonSelFramework\\tests\\screenshot\\"+name+".docx")
